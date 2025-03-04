import streamlit as st
import faiss
import os
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceBgeEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEndpoint
from secret_api_key import huggingface_access_token

print(huggingface_access_token)

def process_input_data(input_type, input_data):
    loader = None
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input file for PDF")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input file for DOCX")
        text = ""
        for para in doc.paragraphs:
            text += para.text
        documents = text
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, UploadedFile):
            text = str(input_data.read().decode('utf-8'))
        else:
            raise ValueError("Invalid file type for .TXT")
    else:
        raise ValueError("Unsupported File/Input type, please upload an acceptable file type")
    
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]
    else:
        texts = text_splitter.split_text(documents)

    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}

    hf_embeddings = HuggingFaceBgeEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs,
    )

    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)

    vector_store = FAISS(
        embedding_function=hf_embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )

    vector_store.add_texts(texts)
    return vector_store

def answer_the_question(vector_store, query):
    llm = HuggingFaceEndpoint(repo_id="meta-llama/Meta-Llama-3-8B-Instruct", token=huggingface_access_token, temperature=0.6)
    qa = RetrievalQA.from_chain_type(llm=llm, retriever=vector_store.as_retriever())
    answer = qa.invoke({"query": query})
    return answer    




def main():
    st.title('RAG QnA App')
    input_type = st.selectbox("Input Type", ["Text","Link", "PDF", "DOCX", "TXT File"])

    if input_type == "Link":
        no_of_links = st.number_input(min_value=1, max_value=20, step=1, label="Select the number of links you want me to refer")
        input_data = []
        for i in range(no_of_links):
            url = st.sidebar.text_input(f"URL {i+1}")
            input_data.append(url)

    elif input_type == "Text":
        input_data = st.text_input("Enter your Text")
    elif input_type == "PDF":
        input_data = st.file_uploader("Upload your PDF file", type=["pdf"])
    elif input_type == "DOCX":
        input_data = st.file_uploader("Upload you DOCX file", type=["docs", "doc"])
    elif input_type == "TXT File":
        input_data = st.file_uploader("Upload you Text file", type=["txt"])


    if st.button("Proceed"):
        vector_store = process_input_data(input_type, input_data)
        st.session_state["vectorstore"] = vector_store
    if "vectorstore" in st.session_state:
        query = st.text_input("Ask your question")
        if st.button ("Submit"):
            answer = answer_the_question(st.session_state["vectorstore"], query)
            st.write(answer)

if __name__ == '__main__':
    main()